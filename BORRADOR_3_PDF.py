import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, ttk
import PyPDF2
import requests
import base64
import os
from pathlib import Path
import tempfile
import re
import cv2
import numpy as np
from PIL import Image
import io
import json

class PDFTextExtractor:
    def __init__(self, root):
        self.root = root
        self.root.title("Extractor de Texto PDF con IA de Fórmulas")
        self.root.geometry("950x750")
        
        # Variable para almacenar el texto extraído
        self.extracted_text = ""
        
        # API Keys
        self.ocr_api_key = "K83967071688957"  # Tu API key para OCR Space
        
        self.setup_ui()
    
    def setup_ui(self):
        # Frame principal
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Configurar expansión
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(1, weight=1)
        main_frame.rowconfigure(6, weight=1)
        
        # Título
        title_label = ttk.Label(main_frame, text="Extractor PDF con IA de Fórmulas", 
                               font=("Arial", 16, "bold"))
        title_label.grid(row=0, column=0, columnspan=4, pady=(0, 20), sticky=tk.W)
        
        # Botón para seleccionar archivo
        self.select_btn = ttk.Button(main_frame, text="Seleccionar PDF", 
                                    command=self.select_file)
        self.select_btn.grid(row=1, column=0, padx=(0, 10), sticky=tk.W)
        
        # Label para mostrar archivo seleccionado
        self.file_label = ttk.Label(main_frame, text="Ningún archivo seleccionado")
        self.file_label.grid(row=1, column=1, columnspan=3, sticky=(tk.W, tk.E))
        
        # Frame para opciones OCR y fórmulas
        options_frame = ttk.LabelFrame(main_frame, text="Opciones de Extracción", padding="5")
        options_frame.grid(row=2, column=0, columnspan=4, sticky=(tk.W, tk.E), pady=(10, 0))
        options_frame.columnconfigure(1, weight=1)
        
        # Checkbox para OCR
        self.use_ocr = tk.BooleanVar()
        self.ocr_checkbox = ttk.Checkbutton(options_frame, text="Usar OCR (PDFs escaneados)", 
                                           variable=self.use_ocr)
        self.ocr_checkbox.grid(row=0, column=0, sticky=tk.W)
        
        # Checkbox para reconocimiento de fórmulas con IA
        self.use_formula_ai = tk.BooleanVar(value=True)
        self.formula_ai_checkbox = ttk.Checkbutton(options_frame, text="IA para fórmulas matemáticas", 
                                                  variable=self.use_formula_ai)
        self.formula_ai_checkbox.grid(row=0, column=1, padx=(20, 0), sticky=tk.W)
        
        # Selector de servicio de fórmulas
        ttk.Label(options_frame, text="Servicio fórmulas:").grid(row=1, column=0, sticky=tk.W, pady=(5, 0))
        self.formula_service = tk.StringVar(value="mathpix")
        formula_combo = ttk.Combobox(options_frame, textvariable=self.formula_service, 
                                    values=["mathpix", "pix2tex", "local_cv"], 
                                    width=15, state="readonly")
        formula_combo.grid(row=1, column=1, sticky=tk.W, padx=(5, 0), pady=(5, 0))
        
        # Selector de motor OCR
        ttk.Label(options_frame, text="Motor OCR:").grid(row=2, column=0, sticky=tk.W, pady=(5, 0))
        self.ocr_engine = tk.StringVar(value="2")
        ocr_combo = ttk.Combobox(options_frame, textvariable=self.ocr_engine, 
                                values=["1 (Básico)", "2 (Avanzado)", "3 (Beta)"], 
                                width=15, state="readonly")
        ocr_combo.grid(row=2, column=1, sticky=tk.W, padx=(5, 0), pady=(5, 0))
        
        # Frame para selección de páginas
        page_frame = ttk.Frame(main_frame)
        page_frame.grid(row=3, column=0, columnspan=4, sticky=(tk.W, tk.E), pady=(10, 0))
        page_frame.columnconfigure(2, weight=1)
        
        # Checkbox para todas las páginas
        self.all_pages = tk.BooleanVar(value=True)
        self.all_pages_checkbox = ttk.Checkbutton(page_frame, text="Todas las páginas", 
                                                 variable=self.all_pages,
                                                 command=self.toggle_page_selection)
        self.all_pages_checkbox.grid(row=0, column=0, sticky=tk.W)
        
        # Label y entry para páginas específicas
        self.pages_label = ttk.Label(page_frame, text="Páginas específicas (ej: 1,3,5-8):")
        self.pages_label.grid(row=0, column=1, padx=(20, 5), sticky=tk.W)
        
        self.pages_entry = ttk.Entry(page_frame, width=20, state="disabled")
        self.pages_entry.grid(row=0, column=2, sticky=tk.W)
        
        # Label informativo sobre total de páginas
        self.total_pages_label = ttk.Label(page_frame, text="", foreground="gray")
        self.total_pages_label.grid(row=0, column=3, padx=(10, 0), sticky=tk.W)

        # Checkbox para separación por páginas
        self.page_separator = tk.BooleanVar(value=True)
        self.page_separator_checkbox = ttk.Checkbutton(
            main_frame, text="Separar texto por páginas", variable=self.page_separator
        )
        self.page_separator_checkbox.grid(row=4, column=0, columnspan=4, sticky=tk.W, pady=(10, 0))
        
        # Frame para información de APIs
        info_frame = ttk.LabelFrame(main_frame, text="Información de APIs", padding="5")
        info_frame.grid(row=5, column=0, columnspan=4, sticky=(tk.W, tk.E), pady=(10, 0))
        
        info_text = ("• Mathpix: Requiere registro gratuito (500 consultas/mes)\n"
                    "• Pix2tex: Servicio gratuito de HuggingFace\n"
                    "• Local CV: Procesamiento local básico")
        ttk.Label(info_frame, text=info_text, font=("Arial", 8)).grid(row=0, column=0, sticky=tk.W)
        
        # Área de texto para mostrar el contenido
        text_frame = ttk.Frame(main_frame)
        text_frame.grid(row=6, column=0, columnspan=4, sticky=(tk.W, tk.E, tk.N, tk.S), pady=(20, 0))
        text_frame.columnconfigure(0, weight=1)
        text_frame.rowconfigure(0, weight=1)
        
        self.text_area = scrolledtext.ScrolledText(text_frame, wrap=tk.WORD, 
                                                  width=80, height=25, font=("Consolas", 10))
        self.text_area.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Frame para botones inferiores
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=7, column=0, columnspan=4, pady=(20, 0))
        
        # Botón para extraer texto
        self.extract_btn = ttk.Button(button_frame, text="Extraer Texto", 
                                     command=self.extract_text, state="disabled")
        self.extract_btn.pack(side=tk.LEFT, padx=(0, 10))
        
        # Botón para configurar APIs
        self.config_btn = ttk.Button(button_frame, text="Configurar APIs", 
                                    command=self.configure_apis)
        self.config_btn.pack(side=tk.LEFT, padx=(0, 10))
        
        # Botón para guardar texto
        self.save_btn = ttk.Button(button_frame, text="Guardar como TXT", 
                                  command=self.save_text, state="disabled")
        self.save_btn.pack(side=tk.LEFT, padx=(0, 10))
        
        # Botón para guardar como Word
        self.save_word_btn = ttk.Button(button_frame, text="Guardar como Word", 
                                       command=self.save_word, state="disabled")
        self.save_word_btn.pack(side=tk.LEFT, padx=(0, 10))
        
        # Botón para limpiar
        self.clear_btn = ttk.Button(button_frame, text="Limpiar", 
                                   command=self.clear_text)
        self.clear_btn.pack(side=tk.LEFT)
        
        # Barra de progreso
        self.progress = ttk.Progressbar(main_frame, mode='indeterminate')
        self.progress.grid(row=8, column=0, columnspan=4, sticky=(tk.W, tk.E), pady=(10, 0))
        
        # Label de estado
        self.status_label = ttk.Label(main_frame, text="Listo")
        self.status_label.grid(row=9, column=0, columnspan=4, pady=(5, 0))
        
        self.selected_file = None
        self.total_pages = 0
        self.mathpix_app_id = ""
        self.mathpix_app_key = ""

    def configure_apis(self):
        """Ventana para configurar APIs de fórmulas"""
        config_window = tk.Toplevel(self.root)
        config_window.title("Configurar APIs")
        config_window.geometry("500x400")
        config_window.transient(self.root)
        config_window.grab_set()
        
        # Frame principal
        frame = ttk.Frame(config_window, padding="10")
        frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        config_window.columnconfigure(0, weight=1)
        config_window.rowconfigure(0, weight=1)
        
        # Mathpix configuration
        mathpix_frame = ttk.LabelFrame(frame, text="Mathpix (Recomendado)", padding="10")
        mathpix_frame.grid(row=0, column=0, sticky=(tk.W, tk.E), pady=(0, 10))
        mathpix_frame.columnconfigure(1, weight=1)
        
        ttk.Label(mathpix_frame, text="1. Regístrate en: https://mathpix.com/").grid(row=0, column=0, columnspan=2, sticky=tk.W)
        ttk.Label(mathpix_frame, text="2. Ve a Dashboard > API").grid(row=1, column=0, columnspan=2, sticky=tk.W)
        
        ttk.Label(mathpix_frame, text="App ID:").grid(row=2, column=0, sticky=tk.W, pady=(10, 5))
        self.app_id_entry = ttk.Entry(mathpix_frame, width=50)
        self.app_id_entry.grid(row=2, column=1, sticky=(tk.W, tk.E), pady=(10, 5))
        self.app_id_entry.insert(0, self.mathpix_app_id)
        
        ttk.Label(mathpix_frame, text="App Key:").grid(row=3, column=0, sticky=tk.W, pady=(0, 5))
        self.app_key_entry = ttk.Entry(mathpix_frame, width=50, show="*")
        self.app_key_entry.grid(row=3, column=1, sticky=(tk.W, tk.E), pady=(0, 5))
        self.app_key_entry.insert(0, self.mathpix_app_key)
        
        # Información adicional
        info_frame = ttk.LabelFrame(frame, text="Servicios Disponibles", padding="10")
        info_frame.grid(row=1, column=0, sticky=(tk.W, tk.E), pady=(0, 10))
        
        info_text = """• Mathpix: Mejor calidad, 500 consultas gratis/mes
• Pix2tex (HuggingFace): Totalmente gratuito, calidad media
• Local CV: Procesamiento básico sin internet

Recomendación: Usar Mathpix para mejores resultados"""
        
        ttk.Label(info_frame, text=info_text, justify=tk.LEFT).grid(row=0, column=0, sticky=tk.W)
        
        # Botones
        button_frame = ttk.Frame(frame)
        button_frame.grid(row=2, column=0, pady=(10, 0))
        
        ttk.Button(button_frame, text="Probar Mathpix", 
                  command=lambda: self.test_mathpix_api()).pack(side=tk.LEFT, padx=(0, 10))
        ttk.Button(button_frame, text="Guardar", 
                  command=lambda: self.save_api_config(config_window)).pack(side=tk.LEFT, padx=(0, 10))
        ttk.Button(button_frame, text="Cancelar", 
                  command=config_window.destroy).pack(side=tk.LEFT)
    
    def test_mathpix_api(self):
        """Probar la configuración de Mathpix"""
        app_id = self.app_id_entry.get().strip()
        app_key = self.app_key_entry.get().strip()
        
        if not app_id or not app_key:
            messagebox.showwarning("Advertencia", "Por favor ingresa App ID y App Key")
            return
        
        try:
            # Crear una imagen de prueba simple con una fórmula
            test_image = Image.new('RGB', (200, 100), color='white')
            # Aquí podrías dibujar una fórmula simple o usar una imagen de prueba
            
            # Convertir a bytes
            img_bytes = io.BytesIO()
            test_image.save(img_bytes, format='PNG')
            img_data = base64.b64encode(img_bytes.getvalue()).decode()
            
            # Hacer petición de prueba
            headers = {
                'app_id': app_id,
                'app_key': app_key,
                'Content-type': 'application/json'
            }
            
            data = {
                'src': f"data:image/png;base64,{img_data}",
                'formats': ['latex_styled']
            }
            
            response = requests.post('https://api.mathpix.com/v3/text', 
                                   headers=headers, json=data, timeout=10)
            
            if response.status_code == 200:
                messagebox.showinfo("Éxito", "¡Conexión con Mathpix exitosa!")
            else:
                messagebox.showerror("Error", f"Error de API: {response.status_code}\n{response.text}")
                
        except Exception as e:
            messagebox.showerror("Error", f"Error probando API: {str(e)}")
    
    def save_api_config(self, window):
        """Guardar configuración de APIs"""
        self.mathpix_app_id = self.app_id_entry.get().strip()
        self.mathpix_app_key = self.app_key_entry.get().strip()
        messagebox.showinfo("Éxito", "Configuración guardada")
        window.destroy()

    def detect_math_regions(self, image):
        """Detectar regiones que pueden contener fórmulas matemáticas usando OpenCV"""
        # Convertir a escala de grises
        gray = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2GRAY)
        
        # Aplicar threshold adaptativo
        thresh = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, 
                                     cv2.THRESH_BINARY_INV, 11, 2)
        
        # Encontrar contornos
        contours, _ = cv2.findContours(thresh, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        
        math_regions = []
        for contour in contours:
            x, y, w, h = cv2.boundingRect(contour)
            # Filtrar por tamaño y forma (características típicas de fórmulas)
            if 20 < w < 500 and 15 < h < 100:
                aspect_ratio = w / h
                if 1.5 < aspect_ratio < 10:  # Fórmulas suelen ser más anchas que altas
                    math_regions.append((x, y, w, h))
        
        return math_regions

    def extract_formula_with_mathpix(self, image_data):
        """Extraer fórmula usando Mathpix API"""
        if not self.mathpix_app_id or not self.mathpix_app_key:
            return None
        
        try:
            headers = {
                'app_id': self.mathpix_app_id,
                'app_key': self.mathpix_app_key,
                'Content-type': 'application/json'
            }
            
            data = {
                'src': f"data:image/png;base64,{image_data}",
                'formats': ['latex_styled', 'text']
            }
            
            response = requests.post('https://api.mathpix.com/v3/text', 
                                   headers=headers, json=data, timeout=15)
            
            if response.status_code == 200:
                result = response.json()
                latex = result.get('latex_styled', '')
                text = result.get('text', '')
                
                if latex and latex.strip():
                    return f"$$${latex}$$$"  # Formato LaTeX
                elif text and text.strip():
                    return text
            
        except Exception as e:
            print(f"Error con Mathpix: {e}")
        
        return None

    def extract_formula_with_pix2tex(self, image_data):
        """Extraer fórmula usando pix2tex de HuggingFace (gratuito)"""
        try:
            API_URL = "https://api-inference.huggingface.co/models/ysharma/pix2tex"
            
            # Decodificar la imagen
            image_bytes = base64.b64decode(image_data)
            
            response = requests.post(API_URL, data=image_bytes, timeout=30)
            
            if response.status_code == 200:
                result = response.json()
                if isinstance(result, list) and len(result) > 0:
                    generated_text = result[0].get('generated_text', '')
                    if generated_text.strip():
                        return f"$$${generated_text}$$$"
            
        except Exception as e:
            print(f"Error con pix2tex: {e}")
        
        return None

    def process_formulas_in_image(self, image):
        """Procesar fórmulas en una imagen"""
        formulas = []
        service = self.formula_service.get()
        
        if service == "local_cv":
            # Procesamiento local básico
            regions = self.detect_math_regions(image)
            for region in regions:
                x, y, w, h = region
                formula_img = image.crop((x, y, x+w, y+h))
                # Aquí podrías implementar reconocimiento básico local
                formulas.append(f"[FÓRMULA DETECTADA EN ({x},{y})]")
        
        elif service in ["mathpix", "pix2tex"]:
            # Detectar regiones con fórmulas
            regions = self.detect_math_regions(image)
            
            for region in regions:
                x, y, w, h = region
                # Extraer región de la fórmula
                formula_img = image.crop((x, y, x+w, y+h))
                
                # Convertir a base64
                img_bytes = io.BytesIO()
                formula_img.save(img_bytes, format='PNG')
                img_data = base64.b64encode(img_bytes.getvalue()).decode()
                
                # Procesar con el servicio seleccionado
                if service == "mathpix":
                    formula_text = self.extract_formula_with_mathpix(img_data)
                else:  # pix2tex
                    formula_text = self.extract_formula_with_pix2tex(img_data)
                
                if formula_text:
                    formulas.append(formula_text)
        
        return formulas

    def fix_text_spacing(self, text):
        """Corregir problemas de espaciado en texto extraído"""
        # Patrón para detectar palabras concatenadas (minúscula seguida de mayúscula)
        text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)
        
        # Patrón para separar números de letras
        text = re.sub(r'([0-9])([A-Za-z])', r'\1 \2', text)
        text = re.sub(r'([A-Za-z])([0-9])', r'\1 \2', text)
        
        # Corregir espacios múltiples
        text = re.sub(r'\s+', ' ', text)
        
        # Mejorar separación de párrafos
        text = re.sub(r'([.!?])\s*([A-Z])', r'\1\n\n\2', text)
        
        return text

    def improve_mathematical_formulas(self, text):
        """Mejorar el reconocimiento de fórmulas matemáticas en texto"""
        # Patrones comunes de fórmulas mal reconocidas
        improvements = [
            # Patrón específico de tu fórmula
            (r'N\s*×\s*\(\s*B\s*-\s*C\s*\)\s*r\s*=\s*N\s*\+\s*V', 
             r'r = N × (B - C) / (N + V)'),
            
            # Mejorar espaciado en fórmulas
            (r'(\w)\s*=\s*(\w)', r'\1 = \2'),
            (r'(\w)\s*\+\s*(\w)', r'\1 + \2'),
            (r'(\w)\s*-\s*(\w)', r'\1 - \2'),
            (r'(\w)\s*×\s*(\w)', r'\1 × \2'),
            (r'(\w)\s*/\s*(\w)', r'\1 / \2'),
            
            # Corregir paréntesis mal espaciados
            (r'\(\s*([^)]+)\s*\)', r'(\1)'),
            
            # Mejorar formato de definiciones de variables
            (r'([A-Z])\s*:\s*([^A-Z\n]+)', r'\1: \2'),
        ]
        
        improved_text = text
        for pattern, replacement in improvements:
            improved_text = re.sub(pattern, replacement, improved_text, 
                                 flags=re.IGNORECASE | re.MULTILINE)
        
        return improved_text

    def post_process_ocr_text(self, text):
        """Post-procesamiento específico para texto OCR"""
        # Correcciones comunes de OCR
        corrections = {
            'CINs': 'CINS',
            'ICATI CADE': 'ICADE',
            '•': '■',
        }
        
        processed_text = text
        for wrong, correct in corrections.items():
            processed_text = processed_text.replace(wrong, correct)
        
        # Mejorar formato de fórmulas
        processed_text = self.improve_mathematical_formulas(processed_text)
        
        return processed_text

    def toggle_page_selection(self):
        """Habilitar/deshabilitar la entrada de páginas específicas"""
        if self.all_pages.get():
            self.pages_entry.config(state="disabled")
            self.pages_entry.delete(0, tk.END)
        else:
            self.pages_entry.config(state="normal")

    def get_total_pages(self, file_path):
        """Obtener el número total de páginas del PDF"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                return len(pdf_reader.pages)
        except:
            return 0

    def parse_page_ranges(self, page_input, total_pages):
        """Parsear la entrada de páginas y devolver una lista de números de página"""
        pages = set()
        
        if not page_input.strip():
            return list(range(1, total_pages + 1))
        
        try:
            parts = page_input.split(',')
            
            for part in parts:
                part = part.strip()
                
                if '-' in part:
                    start, end = part.split('-')
                    start = int(start.strip())
                    end = int(end.strip())
                    
                    if start < 1 or end > total_pages or start > end:
                        raise ValueError(f"Rango inválido: {part}")
                    
                    pages.update(range(start, end + 1))
                else:
                    page_num = int(part)
                    if page_num < 1 or page_num > total_pages:
                        raise ValueError(f"Página fuera de rango: {page_num}")
                    pages.add(page_num)
            
            return sorted(list(pages))
        
        except ValueError as e:
            raise ValueError(f"Formato de páginas inválido: {str(e)}")
        except Exception:
            raise ValueError("Formato de páginas inválido. Use formato: 1,3,5-8")

    def select_file(self):
        """Seleccionar archivo PDF"""
        file_path = filedialog.askopenfilename(
            title="Selecciona un archivo PDF",
            filetypes=[("Archivos PDF", "*.pdf"), ("Todos los archivos", "*.*")]
        )
        
        if file_path:
            self.selected_file = file_path
            filename = os.path.basename(file_path)
            self.file_label.config(text=f"Archivo: {filename}")
            self.extract_btn.config(state="normal")
            self.status_label.config(text=f"Archivo seleccionado: {filename}")
            
            self.total_pages = self.get_total_pages(file_path)
            if self.total_pages > 0:
                self.total_pages_label.config(text=f"(Total: {self.total_pages} páginas)")
            else:
                self.total_pages_label.config(text="(No se pudo determinar el número de páginas)")

    def extract_text_from_pdf(self, file_path, selected_pages=None):
        """Extraer texto usando PyPDF2"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                total_pages = len(pdf_reader.pages)
                
                if selected_pages is None:
                    selected_pages = list(range(1, total_pages + 1))
                
                for page_num in selected_pages:
                    if 1 <= page_num <= total_pages:
                        page = pdf_reader.pages[page_num - 1]
                        page_text = page.extract_text()
                        if page_text.strip():
                            page_text = self.fix_text_spacing(page_text)
                            if self.page_separator.get():
                                text += f"\n--- Página {page_num} ---\n"
                            text += page_text + "\n"
        
        except Exception as e:
            raise Exception(f"Error al leer PDF: {str(e)}")
        
        return text

    def extract_text_with_ocr_and_formulas(self, file_path, selected_pages=None):
        """Extraer texto usando OCR y reconocimiento de fórmulas"""
        try:
            url = 'https://api.ocr.space/parse/image'
            
            # Preparar archivo PDF
            pdf_to_send = file_path
            temp_file = None
            if selected_pages is not None:
                from PyPDF2 import PdfWriter, PdfReader
                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
                writer = PdfWriter()
                with open(file_path, 'rb') as infile:
                    reader = PdfReader(infile)
                    for page_num in selected_pages:
                        if 1 <= page_num <= len(reader.pages):
                            writer.add_page(reader.pages[page_num - 1])
                    writer.write(temp_file)
                temp_file.close()
                pdf_to_send = temp_file.name

          # ...existing code...
            # Extraer texto con OCR
            payload = {
                'apikey': self.ocr_api_key,
                'language': 'spa',
                'OCREngine': self.ocr_engine.get().split()[0],
                'scale': 'true',
                'isTable': 'true',
                'filetype': 'PDF'
            }

            with open(pdf_to_send, 'rb') as f:
                files = {'file': f}
                response = requests.post(url, data=payload, files=files, timeout=60)
            # ...existing code...
            
            # Limpiar archivo temporal si se creó
            if temp_file:
                os.unlink(temp_file.name)
            
            if response.status_code == 200:
                result = response.json()
                if result.get('IsErroredOnProcessing', True):
                    error_msg = result.get('ErrorMessage', ['Error desconocido'])
                    raise Exception(f"Error en OCR: {', '.join(error_msg)}")
                
                # Extraer texto de todas las páginas
                extracted_text = ""
                parsed_results = result.get('ParsedResults', [])
                
                for i, page_result in enumerate(parsed_results):
                    page_text = page_result.get('ParsedText', '')
                    if page_text.strip():
                        page_text = self.fix_text_spacing(page_text)
                        
                        # Procesar fórmulas si está habilitado
                        if self.use_formula_ai.get():
                            # Aquí implementarías el procesamiento de fórmulas
                            # por ahora solo post-procesamos el texto OCR
                            page_text = self.post_process_ocr_text(page_text)
                        
                        if self.page_separator.get() and len(parsed_results) > 1:
                            page_num = selected_pages[i] if selected_pages else i + 1
                            extracted_text += f"\n--- Página {page_num} ---\n"
                        
                        extracted_text += page_text + "\n"
                
                return extracted_text
            else:
                raise Exception(f"Error en la API de OCR: {response.status_code}")
        
        except Exception as e:
            raise Exception(f"Error en OCR: {str(e)}")

    def extract_text(self):
        """Función principal para extraer texto"""
        if not self.selected_file:
            messagebox.showerror("Error", "Por favor selecciona un archivo PDF")
            return
        
        try:
            # Determinar páginas a procesar
            selected_pages = None
            if not self.all_pages.get():
                page_input = self.pages_entry.get().strip()
                if page_input:
                    selected_pages = self.parse_page_ranges(page_input, self.total_pages)
                else:
                    messagebox.showerror("Error", "Por favor especifica las páginas a procesar")
                    return
            
            # Mostrar progreso
            self.progress.start()
            self.status_label.config(text="Extrayendo texto...")
            self.root.update()
            
            # Extraer texto según la opción seleccionada
            if self.use_ocr.get():
                extracted_text = self.extract_text_with_ocr_and_formulas(
                    self.selected_file, selected_pages
                )
            else:
                extracted_text = self.extract_text_from_pdf(
                    self.selected_file, selected_pages
                )
            
            # Mostrar resultado
            if extracted_text.strip():
                self.extracted_text = extracted_text
                self.text_area.delete(1.0, tk.END)
                self.text_area.insert(1.0, extracted_text)
                self.save_btn.config(state="normal")
                self.save_word_btn.config(state="normal")
                
                # Contar páginas procesadas
                pages_processed = len(selected_pages) if selected_pages else self.total_pages
                self.status_label.config(text=f"Texto extraído exitosamente - {pages_processed} páginas procesadas")
            else:
                self.text_area.delete(1.0, tk.END)
                self.text_area.insert(1.0, "No se pudo extraer texto del PDF.\n\nSugerencias:\n" +
                                     "• Si es un PDF escaneado, activa la opción 'Usar OCR'\n" +
                                     "• Verifica que el PDF no esté protegido o dañado\n" +
                                     "• Prueba con diferentes motores de OCR")
                self.status_label.config(text="No se extrajo texto")
                
        except ValueError as ve:
            messagebox.showerror("Error", str(ve))
            self.status_label.config(text="Error en selección de páginas")
        except Exception as e:
            messagebox.showerror("Error", f"Error al extraer texto: {str(e)}")
            self.status_label.config(text="Error en extracción")
        finally:
            self.progress.stop()

    def save_text(self):
        """Guardar texto como archivo TXT"""
        if not self.extracted_text:
            messagebox.showwarning("Advertencia", "No hay texto para guardar")
            return
        
        file_path = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[("Archivos de texto", "*.txt"), ("Todos los archivos", "*.*")],
            title="Guardar texto extraído"
        )
        
        if file_path:
            try:
                with open(file_path, 'w', encoding='utf-8') as file:
                    file.write(self.extracted_text)
                messagebox.showinfo("Éxito", f"Texto guardado en: {file_path}")
                self.status_label.config(text=f"Archivo guardado: {os.path.basename(file_path)}")
            except Exception as e:
                messagebox.showerror("Error", f"Error al guardar archivo: {str(e)}")

    def save_word(self):
        """Guardar texto como archivo Word"""
        if not self.extracted_text:
            messagebox.showwarning("Advertencia", "No hay texto para guardar")
            return
        
        try:
            from docx import Document
            
            file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Archivos Word", "*.docx"), ("Todos los archivos", "*.*")],
                title="Guardar como documento Word"
            )
            
            if file_path:
                doc = Document()
                
                # Agregar título
                if self.selected_file:
                    title = f"Texto extraído de: {os.path.basename(self.selected_file)}"
                    doc.add_heading(title, 0)
                
                # Agregar el texto extraído
                paragraphs = self.extracted_text.split('\n')
                for paragraph in paragraphs:
                    if paragraph.strip():
                        doc.add_paragraph(paragraph)
                
                doc.save(file_path)
                messagebox.showinfo("Éxito", f"Documento Word guardado en: {file_path}")
                self.status_label.config(text=f"Documento guardado: {os.path.basename(file_path)}")
                
        except ImportError:
            messagebox.showerror("Error", "Para guardar como Word necesitas instalar: pip install python-docx")
        except Exception as e:
            messagebox.showerror("Error", f"Error al guardar documento Word: {str(e)}")

    def clear_text(self):
        """Limpiar el área de texto"""
        self.text_area.delete(1.0, tk.END)
        self.extracted_text = ""
        self.save_btn.config(state="disabled")
        self.save_word_btn.config(state="disabled")
        self.status_label.config(text="Texto limpiado")

def main():
    """Función principal"""
    root = tk.Tk()
    app = PDFTextExtractor(root)
    
    # Configurar el comportamiento de cierre
    def on_closing():
        root.quit()
        root.destroy()
    
    root.protocol("WM_DELETE_WINDOW", on_closing)
    
    try:
        root.mainloop()
    except KeyboardInterrupt:
        pass

if __name__ == "__main__":
    main()